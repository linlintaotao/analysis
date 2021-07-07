# *utf-8*
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time
import os


class Reporter:

    def __init__(self, path, configList, rules, cdf):
        self._path = path
        self.configList = configList
        self.rules = rules
        self.cdf = cdf

    def makeReport(self):
        reportDoc = Document()
        reportDoc.add_heading('设备测试数据分析')
        reportDoc.add_paragraph('测试报告生成时间：%s' % time.strftime('%Y%m%d_%H%M%S', time.localtime(time.time())))
        reportDoc.add_paragraph(f'分析数据方式：%s' % self.rules)

        table = reportDoc.add_table(rows=1, cols=7, style="Light Grid Accent 1")
        table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table.style.font.size = Pt(10)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '名称'
        hdr_cells[1].text = '天线距离(m)'
        hdr_cells[2].text = '闰秒(s)'
        hdr_cells[3].text = '固定率'
        hdr_cells[4].text = 'RTK(1σ)'
        hdr_cells[5].text = 'RTK(2σ)'
        hdr_cells[6].text = 'RTK(3σ)'

        for i in range(len(self.configList)):
            row_cells = table.add_row().cells
            row_cells[0].text = self.configList[i][0]
            row_cells[1].text = self.configList[i][2]
            row_cells[2].text = self.configList[i][3]
            row_cells[3].text = self.configList[i][1]
            if i != 0:
                row_cells[4].text = f'%0.2fm' % self.cdf[i - 1]['68%']
                row_cells[5].text = f'%0.2fm' % self.cdf[i - 1]['95%']
                row_cells[6].text = f'%0.2fm' % self.cdf[i - 1]['99.7%']

        reportDoc.add_paragraph(f'测试生成图例文件：')
        self.addPng(reportDoc)
        # report
        reportDoc.save(self._path + '/reportDoc.docx')

    def addPng(self, doc):
        listfile = os.listdir(self._path)
        listfile.sort()
        for fileName in listfile:
            picture = os.path.join(self._path, fileName)
            if picture.endswith('.png'):
                doc.add_picture(picture, Inches(6), Inches(4))
